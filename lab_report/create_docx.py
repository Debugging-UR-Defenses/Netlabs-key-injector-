"""
Generate DOCX file with inventory scripts
Run: pip install python-docx && python create_docx.py
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Title
doc.add_heading('NDG Security+ Lab 2 - Inventory Scripts', 0)

# Part C
doc.add_heading('Part C: Windows System Inventory Script (windowsbox.com)', level=1)
doc.add_paragraph('Save the following as inventory.bat and run it:')

windows_script = '''@echo off
REM Windows System Inventory Script
echo ============================================ > inventory_report.txt
echo WINDOWS SYSTEM INVENTORY REPORT >> inventory_report.txt
echo Generated: %date% %time% >> inventory_report.txt
echo ============================================ >> inventory_report.txt
echo. >> inventory_report.txt
echo --- SYSTEM INFO --- >> inventory_report.txt
systeminfo >> inventory_report.txt
echo. >> inventory_report.txt
echo --- MEMORY INFO --- >> inventory_report.txt
systeminfo | findstr /C:"Physical Memory" >> inventory_report.txt
echo. >> inventory_report.txt
echo --- DISK SPACE --- >> inventory_report.txt
wmic logicaldisk get size,freespace,caption >> inventory_report.txt
echo. >> inventory_report.txt
echo --- CPU INFO --- >> inventory_report.txt
wmic cpu get Name,NumberOfCores,NumberOfLogicalProcessors,LoadPercentage >> inventory_report.txt
echo. >> inventory_report.txt
echo --- USER ACCOUNTS --- >> inventory_report.txt
net user >> inventory_report.txt
echo. >> inventory_report.txt
echo --- SHARED FOLDERS --- >> inventory_report.txt
net share >> inventory_report.txt
echo Inventory saved to inventory_report.txt'''

p = doc.add_paragraph()
run = p.add_run(windows_script)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# Part D
doc.add_heading('Part D: Linux System Inventory Script (linuxbox.com)', level=1)
doc.add_paragraph('Save as inventory.sh and run with: chmod +x inventory.sh && ./inventory.sh')

linux_script = '''#!/bin/bash
OUTPUT_FILE="inventory_report.txt"

echo "============================================" > $OUTPUT_FILE
echo "LINUX SYSTEM INVENTORY REPORT" >> $OUTPUT_FILE
echo "Generated: $(date)" >> $OUTPUT_FILE
echo "============================================" >> $OUTPUT_FILE
echo "" >> $OUTPUT_FILE

echo "--- MEMORY INFO ---" >> $OUTPUT_FILE
free -h >> $OUTPUT_FILE
echo "" >> $OUTPUT_FILE

echo "--- PROCESSOR INFO ---" >> $OUTPUT_FILE
cat /proc/cpuinfo | grep "model name" | head -1 >> $OUTPUT_FILE
echo "Number of CPUs: $(nproc)" >> $OUTPUT_FILE
echo "" >> $OUTPUT_FILE

echo "--- DISK SPACE ---" >> $OUTPUT_FILE
df -h >> $OUTPUT_FILE
echo "" >> $OUTPUT_FILE

echo "--- CPU USAGE ---" >> $OUTPUT_FILE
top -bn1 | head -5 >> $OUTPUT_FILE
echo "" >> $OUTPUT_FILE

echo "--- RUNNING SERVICES ---" >> $OUTPUT_FILE
systemctl list-units --type=service --state=running | head -20 >> $OUTPUT_FILE
echo "" >> $OUTPUT_FILE

echo "--- NETWORK INTERFACES ---" >> $OUTPUT_FILE
ip link show >> $OUTPUT_FILE
echo "" >> $OUTPUT_FILE

echo "Inventory saved to $OUTPUT_FILE"'''

p = doc.add_paragraph()
run = p.add_run(linux_script)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# Part A Scripts
doc.add_heading('Part A: NDG Lab Network Inventory Scripts', level=1)

doc.add_heading('Linux One-Liner (for Ubuntu, Security-Onion, bt/DVL)', level=2)
linux_oneliner = '''H=$(hostname);IP=$(hostname -I 2>/dev/null|awk '{print $1}');MAC=$(ip link show 2>/dev/null|grep -m1 "link/ether"|awk '{print $2}');CPU=$(grep -m1 "model name" /proc/cpuinfo|cut -d: -f2|xargs);RAM=$(free -h|awk '/Mem:/{print $2}');DISK=$(df -h /|awk 'NR==2{print $2}');OS=$(cat /etc/os-release 2>/dev/null|grep "^PRETTY"|cut -d= -f2|tr -d '"');KERN=$(uname -r);echo "=== $H ===" && echo "IP,$IP" && echo "MAC,$MAC" && echo "CPU,$CPU" && echo "RAM,$RAM" && echo "DISK,$DISK" && echo "OS,$OS" && echo "KERN,$KERN" && echo "SSH,$(ssh -V 2>&1|awk '{print $1}')" && echo "SSL,$(openssl version|awk '{print $2}')" && echo "Bash,$(bash --version|head -1|grep -oP '\\d+\\.\\d+\\.\\d+')" && echo "Py,$(python3 --version 2>/dev/null|awk '{print $2}')" && echo "Perl,$(perl -v 2>/dev/null|grep -oP 'v\\d+\\.\\d+' |head -1)" && echo "Apache,$(apache2 -v 2>/dev/null|grep -oP '\\d+\\.\\d+\\.\\d+'||echo -)" && echo "MySQL,$(mysql --version 2>/dev/null|grep -oP '\\d+\\.\\d+\\.\\d+'||echo -)" && echo "PHP,$(php -v 2>/dev/null|grep -oP '\\d+\\.\\d+\\.\\d+'|head -1||echo -)" && echo "sudo,$(sudo --version|head -1|grep -oP '\\d+\\.\\d+\\.\\d+')" && echo "cURL,$(curl --version|head -1|awk '{print $2}')" && lsmod|awk 'NR>1&&NR<12{print "DRV,"$1}' && echo "=== END ==="'''

p = doc.add_paragraph()
run = p.add_run(linux_oneliner)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_heading('Windows PowerShell One-Liner (for WIN12R2, WIN16)', level=2)
ps_oneliner = '''$h=$env:COMPUTERNAME;$ip=(Get-NetIPAddress -AddressFamily IPv4|?{$_.IPAddress -notlike "127.*"}|Select -First 1).IPAddress;$mac=(Get-NetAdapter|Select -First 1).MacAddress;$cpu=(Get-WmiObject Win32_Processor).Name;$ram=[math]::Round((Get-WmiObject Win32_ComputerSystem).TotalPhysicalMemory/1GB,1);$disk=[math]::Round((Get-WmiObject Win32_DiskDrive|Select -First 1).Size/1GB,0);$os=(Get-WmiObject Win32_OperatingSystem).Caption;$ver=(Get-WmiObject Win32_OperatingSystem).Version;Write-Host "=== $h ===";Write-Host "IP,$ip";Write-Host "MAC,$mac";Write-Host "CPU,$cpu";Write-Host "RAM,$ram GB";Write-Host "DISK,$disk GB";Write-Host "OS,$os";Write-Host "VER,$ver";Write-Host "--- SOFTWARE ---";Write-Host "PS,$($PSVersionTable.PSVersion)";Get-WmiObject Win32_Product 2>$null|Select -First 8|%{Write-Host "APP,$($_.Name),$($_.Version)"};Write-Host "--- DRIVERS ---";Get-WmiObject Win32_PnPSignedDriver|Select -First 10|%{Write-Host "DRV,$($_.DeviceName)"};Write-Host "=== END ==="'''

p = doc.add_paragraph()
run = p.add_run(ps_oneliner)
run.font.name = 'Consolas'
run.font.size = Pt(8)

# Usage
doc.add_heading('Usage Instructions', level=1)
doc.add_paragraph('1. Part A (NDG Lab): Copy the appropriate one-liner and paste it into the terminal on each target machine')
doc.add_paragraph('2. Part C (Windows Server): Save the batch script as inventory.bat and run from Command Prompt')
doc.add_paragraph('3. Part D (Linux Server): Save the bash script as inventory.sh, make executable with chmod +x, and run')
doc.add_paragraph('All scripts output to inventory_report.txt for easy review.')

# Save
output_path = 'INVENTORY_SCRIPTS.docx'
doc.save(output_path)
print(f"Created: {output_path}")
